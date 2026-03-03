from pathlib import Path

import docx
import fitz


class DocumentParser:
    @staticmethod
    def _split_header_body(resume_text: str) -> tuple[list[str], list[str]]:
        """
        Split resume text into header lines and body lines.

        Header is the leading block of non-empty lines before the first blank line.
        If there is no blank line, only the first non-empty line is treated as the header.
        """
        raw_lines = resume_text.split('\n')
        header_lines: list[str] = []
        body_lines: list[str] = []

        # Find the first blank line that separates header from body
        blank_found = False
        for line in raw_lines:
            stripped = line.strip()
            if not blank_found:
                if stripped == '':
                    blank_found = True
                else:
                    header_lines.append(stripped)
            else:
                body_lines.append(line)

        # No blank line found: treat only the very first non-empty line as the header
        if not blank_found and len(header_lines) > 1:
            body_lines = header_lines[1:]
            header_lines = header_lines[:1]

        return header_lines, body_lines

    @staticmethod
    def extract_header(resume_text: str) -> str:
        """Extract header (name and contact info) from resume"""
        if not resume_text:
            return ""
        header_lines, _ = DocumentParser._split_header_body(resume_text)
        if not header_lines:
            return ""
        return "[HEADER]\n" + "\n".join(header_lines) + "\n"

    @staticmethod
    def remove_header(resume_text: str) -> str:
        """Remove header from resume text for privacy"""
        if not resume_text:
            return resume_text
        _, body_lines = DocumentParser._split_header_body(resume_text)
        return '\n'.join(body_lines).strip()

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str | None:
        """
        Extract text content from a PDF file using PyMuPDF.

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text content or None if parsing fails
        """
        try:
            pdf_document = fitz.open(file_path)
            text_content = []

            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                text = page.get_text()
                if text.strip():
                    text_content.append(text.strip())

            pdf_document.close()
            return '\n'.join(text_content)

        except Exception as e:
            raise Exception(f"Error parsing PDF file: {str(e)}")

    @staticmethod
    def extract_text(file_path: str) -> str | None:
        """
        Auto-detect file type and extract text from DOCX or PDF.

        Args:
            file_path: Path to the file

        Returns:
            Extracted text content

        Raises:
            ValueError: If file type is not supported
        """
        file_extension = Path(file_path).suffix.lower()

        if file_extension in ['.docx', '.doc']:
            return DocumentParser.extract_text_from_docx(file_path)
        elif file_extension == '.pdf':
            return DocumentParser.extract_text_from_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str | None:
        """
        Extract text content from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text content or None if parsing fails
        """
        try:
            doc = docx.Document(file_path)
            text_content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)

            return '\n'.join(text_content)

        except Exception as e:
            raise Exception(f"Error parsing DOCX file: {str(e)}")

    @staticmethod
    def validate_file(file_path: str, max_size: int) -> bool:
        """
        Validate file existence and size.

        Args:
            file_path: Path to the file
            max_size: Maximum allowed file size in bytes

        Returns:
            True if file is valid, False otherwise
        """
        import os

        if not os.path.exists(file_path):
            return False

        file_size = os.path.getsize(file_path)
        return file_size <= max_size

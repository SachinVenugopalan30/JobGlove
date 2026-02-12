"""Resume parser module for extracting text from DOCX and PDF files."""

import re
from pathlib import Path

import docx
import fitz
import pdfplumber

from utils.logger import app_logger


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file.

    Args:
        file_path: Path to DOCX file

    Returns:
        Extracted text as string

    Raises:
        Exception: If file cannot be read or parsed
    """
    try:
        doc = docx.Document(file_path)
        full_text = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)

        text = '\n'.join(full_text)
        app_logger.info(f"Extracted {len(text)} characters from DOCX file")
        return text

    except Exception as e:
        app_logger.error(f"Error extracting text from DOCX: {e}")
        raise


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using PyMuPDF (fitz).
    Falls back to pdfplumber if fitz fails.

    Args:
        file_path: Path to PDF file

    Returns:
        Extracted text as string

    Raises:
        Exception: If file cannot be read or parsed
    """
    try:
        doc = fitz.open(file_path)
        full_text = []

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            if text.strip():
                full_text.append(text)

        doc.close()
        text = '\n'.join(full_text)

        if len(text.strip()) < 100:
            app_logger.warning("PyMuPDF extracted minimal text, trying pdfplumber fallback")
            text = _extract_text_with_pdfplumber(file_path)

        app_logger.info(f"Extracted {len(text)} characters from PDF file")
        return text

    except Exception as e:
        app_logger.warning(f"PyMuPDF failed: {e}, trying pdfplumber")
        try:
            return _extract_text_with_pdfplumber(file_path)
        except Exception as e2:
            app_logger.error(f"Both PDF parsers failed: {e2}")
            raise


def _extract_text_with_pdfplumber(file_path: str) -> str:
    """
    Fallback PDF extraction using pdfplumber.

    Args:
        file_path: Path to PDF file

    Returns:
        Extracted text as string
    """
    with pdfplumber.open(file_path) as pdf:
        full_text = []
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)

        text = '\n'.join(full_text)
        app_logger.info(f"pdfplumber extracted {len(text)} characters")
        return text


def extract_text(file_path: str) -> str:
    """
    Auto-detect file type and extract text.

    Args:
        file_path: Path to resume file (PDF or DOCX)

    Returns:
        Extracted text as string

    Raises:
        ValueError: If file type is not supported
        Exception: If file cannot be read or parsed
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    extension = path.suffix.lower()

    if extension == '.docx':
        return extract_text_from_docx(file_path)
    elif extension == '.doc':
        raise ValueError("DOC format not yet supported. Please convert to DOCX or PDF.")
    elif extension == '.pdf':
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {extension}. Supported types: .docx, .pdf")


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.

    Args:
        text: Raw extracted text

    Returns:
        Cleaned text

    Operations:
    - Remove excessive whitespace
    - Normalize line breaks
    - Remove special characters (preserve alphanumeric, common punctuation)
    - Strip leading/trailing whitespace
    """
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\t+', ' ', text)
    text = text.strip()

    return text

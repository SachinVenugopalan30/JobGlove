"""Tests for NLP-related API routes."""

import pytest
import os
import tempfile
from io import BytesIO
from docx import Document


@pytest.fixture
def sample_resume_docx():
    """Create a sample DOCX resume for testing."""
    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("john.doe@example.com")
    doc.add_paragraph("(555) 123-4567")
    doc.add_paragraph("")
    doc.add_heading("EXPERIENCE", level=1)
    doc.add_paragraph("Senior Software Engineer at Tech Corp (2019-2024)")
    doc.add_paragraph("- Developed web applications using Python and React")
    doc.add_paragraph("- Led team of 5 developers")
    doc.add_paragraph("- Increased system performance by 40%")
    doc.add_paragraph("")
    doc.add_heading("EDUCATION", level=1)
    doc.add_paragraph("Bachelor's in Computer Science")
    doc.add_paragraph("")
    doc.add_heading("SKILLS", level=1)
    doc.add_paragraph("Python, JavaScript, React, Docker, AWS, SQL")

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()

    yield temp_file.name

    if os.path.exists(temp_file.name):
        os.remove(temp_file.name)


@pytest.fixture
def sample_job_description():
    """Sample job description for testing."""
    return """
    We are looking for a Senior Software Engineer with 5+ years of experience.

    Requirements:
    - Strong experience with Python and React
    - Experience with AWS and Docker
    - Leadership experience
    - Bachelor's degree in Computer Science

    Responsibilities:
    - Build scalable web applications
    - Lead engineering teams
    - Implement best practices
    """


def test_check_apis(client):
    """Test API availability check endpoint."""
    response = client.get('/api/check-apis')
    assert response.status_code == 200

    data = response.get_json()
    assert 'openai' in data
    assert 'gemini' in data
    assert 'claude' in data
    assert isinstance(data['openai'], bool)


def test_upload_resume_docx(client, sample_resume_docx):
    """Test uploading a DOCX resume."""
    with open(sample_resume_docx, 'rb') as f:
        data = {
            'file': (f, 'test_resume.docx')
        }
        response = client.post(
            '/api/upload-resume',
            data=data,
            content_type='multipart/form-data'
        )

    assert response.status_code == 200
    result = response.get_json()
    assert 'file_path' in result
    assert 'filename' in result

    if 'file_path' in result and os.path.exists(result['file_path']):
        os.remove(result['file_path'])


def test_upload_resume_invalid_type(client):
    """Test uploading an invalid file type."""
    data = {
        'file': (BytesIO(b'test content'), 'test.txt')
    }
    response = client.post(
        '/api/upload-resume',
        data=data,
        content_type='multipart/form-data'
    )

    assert response.status_code == 400
    result = response.get_json()
    assert 'error' in result


def test_analyze_resume(client, sample_resume_docx, sample_job_description):
    """Test analyzing a resume with local NLP."""
    with open(sample_resume_docx, 'rb') as f:
        upload_response = client.post(
            '/api/upload-resume',
            data={'file': (f, 'test_resume.docx')},
            content_type='multipart/form-data'
        )

    assert upload_response.status_code == 200
    upload_data = upload_response.get_json()
    file_path = upload_data['file_path']

    try:
        analyze_response = client.post(
            '/api/analyze-resume',
            json={
                'file_path': file_path,
                'job_description': sample_job_description
            }
        )

        assert analyze_response.status_code == 200
        result = analyze_response.get_json()

        assert 'extracted_data' in result
        assert 'score' in result

        extracted_data = result['extracted_data']
        assert 'name' in extracted_data
        assert 'email' in extracted_data
        assert 'phone' in extracted_data

        score = result['score']
        assert 'total_score' in score
        assert 'keyword_match_score' in score
        assert 'keyword_match_details' in score
        assert 'relevance_score' in score
        assert 'ats_score' in score
        assert 'quality_score' in score
        assert 'recommendations' in score

        assert 0 <= score['total_score'] <= 100
        assert isinstance(score['recommendations'], list)

    finally:
        if os.path.exists(file_path):
            os.remove(file_path)


def test_analyze_resume_missing_fields(client):
    """Test analyze endpoint with missing fields."""
    response = client.post(
        '/api/analyze-resume',
        json={'file_path': '/fake/path.docx'}
    )

    assert response.status_code == 400
    result = response.get_json()
    assert 'error' in result


def test_analyze_resume_file_not_found(client, sample_job_description):
    """Test analyze endpoint with non-existent file."""
    response = client.post(
        '/api/analyze-resume',
        json={
            'file_path': '/nonexistent/file.docx',
            'job_description': sample_job_description
        }
    )

    assert response.status_code == 404
    result = response.get_json()
    assert 'error' in result
